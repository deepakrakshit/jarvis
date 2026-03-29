"""DOCX parser for hybrid document intelligence pipeline.

Responsibilities:
- Extract paragraph and table text via python-docx
- Extract embedded image payloads for downstream vision processing
"""

from __future__ import annotations

import importlib.util
import logging
import sys
from typing import Any

from services.document.models import PageContent, RawExtractionResult, TableData
from services.document.parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

_MAX_VISION_IMAGES = 12


class DocxParser(BaseParser):
    """Parser for Microsoft Word (.docx) documents."""

    def parse(self, file_path: str) -> RawExtractionResult:
        try:
            return self._parse_internal(file_path)
        except Exception as exc:
            logger.exception("DOCX parsing failed for %s", file_path)
            return RawExtractionResult(
                text="",
                pages=[],
                tables=[],
                metadata={"file_path": file_path},
                source_type="docx",
                file_path=file_path,
                error=f"DOCX parsing failed: {exc}",
            )

    def _parse_internal(self, file_path: str) -> RawExtractionResult:
        try:
            import docx
        except ImportError as exc:
            spec = importlib.util.find_spec("docx")
            if spec is None:
                error_text = (
                    "python-docx is not installed in the active interpreter. "
                    f"Interpreter: {sys.executable}. "
                    "Install with: python -m pip install python-docx"
                )
            else:
                error_text = (
                    "python-docx appears installed but failed to import in the active interpreter. "
                    f"Interpreter: {sys.executable}. "
                    f"Original import error: {exc}"
                )

            return RawExtractionResult(
                text="",
                pages=[],
                tables=[],
                metadata={"file_path": file_path},
                source_type="docx",
                file_path=file_path,
                error=error_text,
            )

        document = docx.Document(file_path)
        metadata = self._extract_metadata(document, file_path)

        text_parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style_name:
                level = self._heading_level(style_name)
                text_parts.append(f"{'#' * level} {text}")
            else:
                text_parts.append(text)

        tables = self._extract_tables(document)
        embedded_images = self._extract_images(document)
        combined_text = "\n\n".join(text_parts)

        has_images = bool(embedded_images)
        low_text = len(combined_text) < 100

        metadata.update(
            {
                "source_type": "docx",
                "has_images": has_images,
                "is_scanned": bool(has_images and low_text),
                "vision_images": embedded_images,
                "ocr_images": embedded_images if low_text else [],
                "vision_image_count": len(embedded_images),
                "ocr_image_count": len(embedded_images if low_text else []),
            }
        )

        pages = [
            PageContent(
                page_number=1,
                text=combined_text,
                has_images=has_images,
            )
        ]

        return RawExtractionResult(
            text=combined_text,
            pages=pages,
            tables=tables,
            metadata=metadata,
            source_type="docx",
            file_path=file_path,
        )

    @staticmethod
    def _extract_metadata(document: Any, file_path: str) -> dict[str, Any]:
        meta: dict[str, Any] = {"file_path": file_path}
        try:
            props = document.core_properties
            if props.title:
                meta["title"] = str(props.title).strip()
            if props.author:
                meta["author"] = str(props.author).strip()
            if props.subject:
                meta["subject"] = str(props.subject).strip()
            if props.created:
                meta["created"] = str(props.created)
            if props.modified:
                meta["modified"] = str(props.modified)
        except Exception:
            pass

        meta["paragraph_count"] = len(document.paragraphs)
        meta["table_count"] = len(document.tables)
        return meta

    @staticmethod
    def _extract_tables(document: Any) -> list[TableData]:
        tables: list[TableData] = []

        for table in document.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if not rows_data:
                continue

            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []
            data_rows = [
                row for row in data_rows
                if any(cell.strip() for cell in row)
            ]

            if headers or data_rows:
                tables.append(
                    TableData(
                        page_number=1,
                        headers=headers,
                        rows=data_rows,
                    )
                )

        return tables

    @staticmethod
    def _extract_images(document: Any) -> list[dict[str, Any]]:
        images: list[dict[str, Any]] = []
        try:
            for rel_id, rel in document.part.rels.items():
                rel_type = str(getattr(rel, "reltype", "")).lower()
                if "image" not in rel_type:
                    continue

                target_part = getattr(rel, "target_part", None)
                if target_part is None:
                    continue

                image_bytes = getattr(target_part, "blob", None)
                if not isinstance(image_bytes, (bytes, bytearray)) or not image_bytes:
                    continue

                content_type = str(getattr(target_part, "content_type", "image/png") or "image/png")
                if not content_type.startswith("image/"):
                    content_type = "image/png"

                images.append(
                    {
                        "source": f"docx_embedded_{rel_id}",
                        "mime_type": content_type,
                        "bytes": bytes(image_bytes),
                    }
                )
                if len(images) >= _MAX_VISION_IMAGES:
                    break
        except Exception as exc:
            logger.debug("DOCX image extraction failed: %s", exc)

        return images

    @staticmethod
    def _heading_level(style_name: str) -> int:
        import re

        match = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
        if match:
            return min(int(match.group(1)), 6)
        if "title" in style_name:
            return 1
        return 1
